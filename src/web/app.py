"""
src/web/app.py — v4 웹 UI (Day 122, Step 11, 덮어쓰기)

FastAPI + Jinja2 템플릿 기반 최소 MVP.
  - 7-Phase 워크플로우 UI
  - 폼 + 버튼 + 상태 표시 (실시간 기능 없음, v4.1로 이연)
  - 인증 없음 (단일 사용자 로컬 전용)

라우트:
  GET  /                            — 새 프로젝트 시작 페이지
  POST /projects                    — 새 프로젝트 생성 (Phase 0.5 실행)
  GET  /projects                    — 전체 프로젝트 리스트
  GET  /project/{pid}               — 프로젝트 상태 대시보드
  POST /project/{pid}/phase-1       — Phase 1 실행
  POST /project/{pid}/phase-2       — Phase 2 실행
  POST /project/{pid}/phase-3       — Phase 3 실행
  POST /project/{pid}/phase-4       — Phase 4 실행 (옵션)
  GET  /project/{pid}/phase-5       — 피드백 입력 페이지
  POST /project/{pid}/phase-5/feedback — 피드백 적용
  POST /project/{pid}/phase-5/confirm  — 최종 확정
  GET  /project/{pid}/phase-6       — 트랙 선택 페이지
  POST /project/{pid}/phase-6       — 결정 제출
  GET  /project/{pid}/phase-7       — Phase 7 대시보드
  POST /project/{pid}/phase-7/start — Phase 7 진입
  POST /project/{pid}/phase-7/approval — 승인 결정
  POST /project/{pid}/phase-7/packet   — 패킷 생성
  POST /project/{pid}/phase-7/execution-result — 실행 결과 입력
  POST /project/{pid}/phase-7/verification     — 검증
  POST /project/{pid}/phase-7/finalize         — 최종 요약
"""
from __future__ import annotations

import os
from pathlib import Path

# ---------------------------------------------------------------------------
# .env 자동 로드 (Step 13 후속 — uvicorn 기동 시 API 키 주입)
# 반드시 os.environ.get(...) 호출 전에 실행되어야 함.
# python-dotenv 미설치 환경에서도 조용히 건너뜀.
# ---------------------------------------------------------------------------
try:
    from dotenv import load_dotenv
    # 프로젝트 루트의 .env를 찾아 로드 (이미 환경변수로 설정된 값은 덮어쓰지 않음)
    _ROOT = Path(__file__).resolve().parent.parent.parent
    load_dotenv(dotenv_path=_ROOT / ".env", override=False)
except ImportError:
    pass

from fastapi import FastAPI, Form, Request, UploadFile, File
from fastapi.responses import HTMLResponse, RedirectResponse, Response
from fastapi.staticfiles import StaticFiles
from fastapi.templating import Jinja2Templates

from src.web.handlers import (
    handle_phase_0_5, handle_phase_1, handle_phase_2,
    handle_phase_3, handle_phase_4,
    handle_phase_5_feedback, handle_phase_5_confirm,
    handle_phase_6,
    handle_phase_7_start, handle_phase_7_approval, handle_phase_7_packet,
    handle_phase_7_execution_result, handle_phase_7_verification,
    handle_phase_7_finalize,
    get_project_status, list_all_projects,
)


# ---------------------------------------------------------------------------
# 앱 구성
# ---------------------------------------------------------------------------

DB_PATH = os.environ.get("ORCHESTRATOR_DB_PATH", "orchestrator_v1.db")
BASE_DIR = Path(__file__).parent
TEMPLATES_DIR = BASE_DIR / "templates"
STATIC_DIR = BASE_DIR / "static"

app = FastAPI(title="AI Orchestrator v4", version="4.0.0")

if STATIC_DIR.exists():
    app.mount("/static", StaticFiles(directory=str(STATIC_DIR)), name="static")

templates = Jinja2Templates(directory=str(TEMPLATES_DIR))


# ---------------------------------------------------------------------------
# 루트 — 새 프로젝트
# ---------------------------------------------------------------------------

@app.get("/", response_class=HTMLResponse)
async def index(request: Request):
    return templates.TemplateResponse(request, "index.html", {
            "error": None},
    )


@app.post("/projects")
async def create_project(
    request: Request,
    raw_input: str = Form(""),
    project_type: str = Form("doc_generation"),
    template_file: UploadFile = File(None),
    context_files: list[UploadFile] = File(None),
):
    # 프로젝트 종류 검증
    if project_type not in ("doc_generation", "app_dev"):
        project_type = "doc_generation"

    # 양식 파일 파싱 (문서 생성 모드 · 단일 · 선택)
    template_text = ""
    if template_file and template_file.filename:
        try:
            template_text = await _extract_template_text(template_file)
        except Exception as e:
            return templates.TemplateResponse(request, "index.html", {
                "error": f"양식 파일을 읽을 수 없습니다: {e}",
            })

    # 기획문서 묶음 파싱 (앱개발 모드 · 다중)
    referenced_context = None
    if project_type == "app_dev" and context_files:
        try:
            referenced_context = await _extract_referenced_context(context_files)
        except Exception as e:
            return templates.TemplateResponse(request, "index.html", {
                "error": f"기획문서 처리 실패: {e}",
            })

    result = handle_phase_0_5(
        raw_input,
        DB_PATH,
        template_text=template_text,
        project_type=project_type,
        referenced_context=referenced_context,
    )
    if not result["ok"]:
        return templates.TemplateResponse(request, "index.html", {
            "error": result.get("error")},
        )
    return RedirectResponse(
        f"/project/{result['project_id']}", status_code=303,
    )


@app.get("/projects", response_class=HTMLResponse)
async def projects_list(request: Request):
    result = list_all_projects(DB_PATH)
    return templates.TemplateResponse(request, "projects_list.html", {
            "projects": result["projects"], "error": result.get("error")},
    )


# ---------------------------------------------------------------------------
# 프로젝트 상태 대시보드
# ---------------------------------------------------------------------------

@app.get("/project/{project_id}", response_class=HTMLResponse)
async def project_status(project_id: str, request: Request):
    status = get_project_status(project_id, DB_PATH)
    if not status["ok"]:
        return templates.TemplateResponse(request, "index.html", {
            "error": status.get("error")},
        )
    return templates.TemplateResponse(request, "project_status.html", {
            "project": status["project"],
            "runs": status["runs"],
            "artifact": status["runs"][0] if status["runs"] else {},
        },
    )


# ---------------------------------------------------------------------------
# Phase 1 — 서브주제 분해
# ---------------------------------------------------------------------------

@app.post("/project/{project_id}/phase-1")
async def phase_1(project_id: str):
    handle_phase_1(project_id, DB_PATH)
    return RedirectResponse(f"/project/{project_id}", status_code=303)


# ---------------------------------------------------------------------------
# Phase 2 — 병렬 리서치 (동기 + 긴 타임아웃)
# ---------------------------------------------------------------------------

@app.post("/project/{project_id}/phase-2")
async def phase_2(project_id: str, deep_research: str = Form(default="")):
    mode = "deep_research" if deep_research else "web_search"
    handle_phase_2(project_id, DB_PATH, mode=mode)
    return RedirectResponse(f"/project/{project_id}", status_code=303)


# ---------------------------------------------------------------------------
# Phase 3 — 2문서 합성
# ---------------------------------------------------------------------------

@app.post("/project/{project_id}/phase-3")
async def phase_3(project_id: str):
    handle_phase_3(project_id, DB_PATH)
    return RedirectResponse(f"/project/{project_id}", status_code=303)


# ---------------------------------------------------------------------------
# Phase 4 — 3감사관 + 통합
# ---------------------------------------------------------------------------

@app.post("/project/{project_id}/phase-4")
async def phase_4(project_id: str):
    handle_phase_4(project_id, DB_PATH)
    return RedirectResponse(f"/project/{project_id}", status_code=303)


# ---------------------------------------------------------------------------
# Phase 5 — 피드백 루프
# ---------------------------------------------------------------------------

@app.get("/project/{project_id}/phase-5", response_class=HTMLResponse)
async def phase_5_page(project_id: str, request: Request):
    status = get_project_status(project_id, DB_PATH)
    if not status["ok"]:
        return RedirectResponse("/", status_code=303)

    artifact = status["runs"][0] if status["runs"] else {}
    return templates.TemplateResponse(request, "phase_5.html", {
            "project": status["project"],
            "target_doc": artifact.get("target_doc"),
            "doc_versions": artifact.get("doc_versions") or [],
        },
    )


@app.post("/project/{project_id}/phase-5/feedback")
async def phase_5_feedback(project_id: str, user_feedback: str = Form(...)):
    handle_phase_5_feedback(project_id, DB_PATH, user_feedback)
    return RedirectResponse(f"/project/{project_id}/phase-5", status_code=303)


@app.post("/project/{project_id}/phase-5/confirm")
async def phase_5_confirm(project_id: str):
    handle_phase_5_confirm(project_id, DB_PATH)
    return RedirectResponse(f"/project/{project_id}/phase-6", status_code=303)


# ---------------------------------------------------------------------------
# Phase 6 — 트랙 선택
# ---------------------------------------------------------------------------

@app.get("/project/{project_id}/phase-6", response_class=HTMLResponse)
async def phase_6_page(project_id: str, request: Request):
    status = get_project_status(project_id, DB_PATH)
    if not status["ok"]:
        return RedirectResponse("/", status_code=303)

    artifact = status["runs"][0] if status["runs"] else {}
    return templates.TemplateResponse(request, "phase_6.html", {
            "project": status["project"],
            "target_doc": artifact.get("target_doc"),
        },
    )


@app.post("/project/{project_id}/phase-6")
async def phase_6_decide(project_id: str, decision: str = Form(...)):
    result = handle_phase_6(project_id, DB_PATH, decision)
    if result["ok"] and result["decision"] == "app_dev":
        return RedirectResponse(f"/project/{project_id}/phase-7", status_code=303)
    return RedirectResponse(f"/project/{project_id}", status_code=303)


# ---------------------------------------------------------------------------
# Phase 7 — 앱개발
# ---------------------------------------------------------------------------

@app.get("/project/{project_id}/phase-7", response_class=HTMLResponse)
async def phase_7_page(project_id: str, request: Request):
    status = get_project_status(project_id, DB_PATH)
    if not status["ok"]:
        return RedirectResponse("/", status_code=303)

    # Phase 7 실행된 artifact 찾기 (phase=phase_7인 가장 최근)
    phase_7_runs = [r for r in status["runs"] if r.get("phase") == "phase_7"]
    phase_7_artifact = phase_7_runs[0] if phase_7_runs else {}

    return templates.TemplateResponse(request, "phase_7.html", {
            "project": status["project"],
            "all_runs": status["runs"],
            "phase_7_artifact": phase_7_artifact,
        },
    )


@app.post("/project/{project_id}/phase-7/start")
async def phase_7_start(project_id: str):
    handle_phase_7_start(project_id, DB_PATH)
    return RedirectResponse(f"/project/{project_id}/phase-7", status_code=303)


@app.post("/project/{project_id}/phase-7/approval")
async def phase_7_approval(
    project_id: str, run_id: str = Form(...), decision: str = Form(...),
):
    handle_phase_7_approval(run_id, DB_PATH, decision)
    return RedirectResponse(f"/project/{project_id}/phase-7", status_code=303)


@app.post("/project/{project_id}/phase-7/packet")
async def phase_7_packet(project_id: str, run_id: str = Form(...)):
    handle_phase_7_packet(run_id, DB_PATH, ".")
    return RedirectResponse(f"/project/{project_id}/phase-7", status_code=303)


@app.post("/project/{project_id}/phase-7/execution-result")
async def phase_7_execution_result(
    project_id: str,
    run_id: str = Form(...),
    changed_files: str = Form(""),
    test_results: str = Form(""),
    run_log: str = Form(""),
):
    files_list = [f.strip() for f in changed_files.split(",") if f.strip()]
    handle_phase_7_execution_result(
        run_id, DB_PATH, files_list, test_results, run_log,
    )
    return RedirectResponse(f"/project/{project_id}/phase-7", status_code=303)


@app.post("/project/{project_id}/phase-7/verification")
async def phase_7_verification(project_id: str, run_id: str = Form(...)):
    handle_phase_7_verification(run_id, DB_PATH)
    return RedirectResponse(f"/project/{project_id}/phase-7", status_code=303)


@app.post("/project/{project_id}/phase-7/finalize")
async def phase_7_finalize(project_id: str, run_id: str = Form(...)):
    handle_phase_7_finalize(run_id, DB_PATH)
    return RedirectResponse(f"/project/{project_id}/phase-7", status_code=303)


# ---------------------------------------------------------------------------
# 다운로드 — 보고서 파일로 받기
# ---------------------------------------------------------------------------

@app.get("/project/{project_id}/download/{format}")
async def download_report(project_id: str, format: str):
    """
    프로젝트의 target_doc을 파일로 다운로드.

    format:
      - md   : 마크다운
      - txt  : 일반 텍스트 (마크다운 마커 그대로)
      - docx : 워드 문서 (python-docx 필요)
    """
    if format not in ("md", "txt", "docx"):
        return Response("Invalid format. Use 'md', 'txt', or 'docx'.", status_code=400)

    status = get_project_status(project_id, DB_PATH)
    if not status["ok"]:
        return Response("Project not found", status_code=404)

    artifact = status["runs"][0] if status["runs"] else {}
    target_doc = artifact.get("target_doc")
    if not target_doc or not target_doc.get("document"):
        return Response("No document available yet.", status_code=404)

    document_text = target_doc["document"]

    # 파일명 — 프로젝트 제목 사용 (한글 안전 처리)
    import urllib.parse
    title = (status["project"].get("title") or "report").strip()

    # ASCII 안전 버전 (한글/특수문자 → _)
    ascii_safe = "".join(
        c if c.isascii() and (c.isalnum() or c in "-_.") else "_"
        for c in title
    )[:80].strip("_") or "report"

    # UTF-8 인코딩된 원본 (RFC 5987)
    utf8_encoded = urllib.parse.quote(title[:80])

    filename_ascii = f"{ascii_safe}.{format}"
    filename_utf8 = f"{utf8_encoded}.{format}"

    # docx 형식: python-docx로 변환
    if format == "docx":
        try:
            from io import BytesIO
            from docx import Document
            from docx.shared import Pt
        except ImportError:
            return Response(
                "python-docx 라이브러리가 설치되어 있지 않습니다.\n"
                "pip install python-docx",
                status_code=500,
                media_type="text/plain; charset=utf-8",
            )

        doc = Document()
        # 기본 폰트
        try:
            doc.styles["Normal"].font.name = "맑은 고딕"
            doc.styles["Normal"].font.size = Pt(11)
        except Exception:
            pass

        # 마크다운 → docx 변환 (헬퍼 함수에서 처리)
        _markdown_to_docx(doc, document_text)

        buf = BytesIO()
        doc.save(buf)
        buf.seek(0)

        return Response(
            content=buf.read(),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={
                "Content-Disposition": f"attachment; filename=\"{filename_ascii}\"; filename*=UTF-8''{filename_utf8}",
            },
        )

    # md / txt 형식
    media_type = "text/markdown" if format == "md" else "text/plain"

    return Response(
        content=document_text,
        media_type=f"{media_type}; charset=utf-8",
        headers={
            "Content-Disposition": f"attachment; filename=\"{filename_ascii}\"; filename*=UTF-8''{filename_utf8}",
        },
    )


# ---------------------------------------------------------------------------
# 양식 파일 → 텍스트 추출
# ---------------------------------------------------------------------------

async def _extract_template_text(upload: "UploadFile") -> str:
    """
    업로드된 양식 파일을 텍스트로 변환.

    지원 형식:
      - .docx : python-docx로 단락별 텍스트 추출
      - .md, .txt, 기타 텍스트 : 그대로 디코드 (UTF-8)

    크기 제한: 200KB (LLM 토큰 폭발 방지)
    """
    MAX_BYTES = 200 * 1024  # 200KB

    # 비동기 read
    content = await upload.read()
    if not content:
        return ""
    if len(content) > MAX_BYTES:
        raise ValueError(
            f"양식 파일이 너무 큽니다 ({len(content)//1024}KB). "
            f"최대 {MAX_BYTES//1024}KB까지 허용됩니다."
        )

    filename = (upload.filename or "").lower()

    # docx
    if filename.endswith(".docx"):
        try:
            from io import BytesIO
            from docx import Document
        except ImportError:
            raise ValueError("python-docx 라이브러리 필요 (pip install python-docx)")

        try:
            doc = Document(BytesIO(content))
        except Exception as e:
            raise ValueError(f".docx 파일 읽기 실패: {e}")

        lines = []
        for para in doc.paragraphs:
            text = (para.text or "").strip()
            if text:
                lines.append(text)
        # 표도 추출
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells)
                if row_text.strip(" |"):
                    lines.append(row_text)

        return "\n".join(lines).strip()

    # md / txt / 그 외 — 텍스트로 디코드
    for encoding in ("utf-8", "utf-8-sig", "cp949", "euc-kr"):
        try:
            return content.decode(encoding).strip()
        except UnicodeDecodeError:
            continue

    # 모두 실패
    raise ValueError("파일 인코딩을 인식할 수 없습니다 (UTF-8/CP949/EUC-KR 시도)")


async def _extract_referenced_context(uploads: list) -> dict:
    """
    앱개발 모드에서 업로드된 기획문서 묶음을 영구 저장 가능한 형태로 변환.

    각 파일을 _extract_template_text()와 동일하게 파싱한 뒤,
    파일별 메타데이터와 함께 하나의 dict로 묶음.

    Returns:
      {
        "files": [
          {
            "filename": "CLAUDE.md",
            "role": "context",       # 향후 자동 분류 가능
            "content": "...",
            "size_bytes": 12345
          },
          ...
        ],
        "uploaded_at": "2026-04-27T..."
      }

    제한:
      - 파일별 200KB (각 _extract_template_text가 검증)
      - 전체 합계 1MB (이 함수에서 검증)
      - 빈 업로드 (filename 없음)는 무시
    """
    from datetime import datetime, timezone

    MAX_TOTAL_BYTES = 1024 * 1024  # 1MB

    files = []
    total_bytes = 0

    for upload in uploads:
        # 빈 슬롯 무시 (사용자가 input을 비웠을 때 FastAPI가 빈 UploadFile을 줄 수 있음)
        if not upload or not upload.filename:
            continue

        # 파일 파싱 (재사용 — 200KB 검증·.docx/.md/.txt 처리·인코딩 자동감지 모두 동일)
        content = await _extract_template_text(upload)
        if not content:
            continue

        size = len(content.encode("utf-8"))
        total_bytes += size

        if total_bytes > MAX_TOTAL_BYTES:
            raise ValueError(
                f"기획문서 묶음 합계가 너무 큽니다 ({total_bytes//1024}KB). "
                f"최대 {MAX_TOTAL_BYTES//1024}KB까지 허용됩니다."
            )

        files.append({
            "filename": upload.filename,
            "role": "context",
            "content": content,
            "size_bytes": size,
        })

    if not files:
        return None  # 빈 업로드 → None 처리 (DB에 NULL)

    return {
        "files": files,
        "uploaded_at": datetime.now(timezone.utc).isoformat(),
    }


# ---------------------------------------------------------------------------
# Markdown → docx 변환 헬퍼
# ---------------------------------------------------------------------------

def _markdown_to_docx(doc, text: str) -> None:
    """
    마크다운 텍스트를 docx Document에 추가.

    지원하는 형식:
      - # / ## / ### / #### 헤딩
      - - / * 글머리표
      - 1. 2. 3. 번호 매기기
      - **bold** / *italic*
      - [text](url) 링크 (텍스트만, 하이퍼링크 X)
      - 마크다운 표 (| col1 | col2 |)
      - --- 구분선 (빈 줄)
      - ``` 코드 블록 (고정폭 폰트 단락)

    제한:
      - 중첩 리스트 미지원 (1단 평면)
      - 이미지 미지원
      - 인라인 코드 ` ` 는 일반 텍스트로
    """
    lines = text.split("\n")
    i = 0
    in_code_block = False
    code_buffer = []

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # ----- 코드 블록 -----
        if stripped.startswith("```"):
            if in_code_block:
                # 닫는 ```
                _add_code_block(doc, "\n".join(code_buffer))
                code_buffer = []
                in_code_block = False
            else:
                # 여는 ```
                in_code_block = True
            i += 1
            continue

        if in_code_block:
            code_buffer.append(line)
            i += 1
            continue

        # ----- 빈 줄 -----
        if not stripped:
            doc.add_paragraph("")
            i += 1
            continue

        # ----- 표 (헤더 + 구분 + 데이터) -----
        if stripped.startswith("|") and i + 1 < len(lines):
            next_line = lines[i + 1].strip()
            # 표 구분선 패턴: |---|---|
            if next_line.startswith("|") and "-" in next_line:
                table_lines = []
                j = i
                while j < len(lines) and lines[j].strip().startswith("|"):
                    table_lines.append(lines[j])
                    j += 1
                _add_markdown_table(doc, table_lines)
                i = j
                continue

        # ----- 헤딩 -----
        if stripped.startswith("# "):
            doc.add_heading(_strip_inline_marks(stripped[2:].strip()), level=1)
            i += 1
            continue
        if stripped.startswith("## "):
            doc.add_heading(_strip_inline_marks(stripped[3:].strip()), level=2)
            i += 1
            continue
        if stripped.startswith("### "):
            doc.add_heading(_strip_inline_marks(stripped[4:].strip()), level=3)
            i += 1
            continue
        if stripped.startswith("#### "):
            doc.add_heading(_strip_inline_marks(stripped[5:].strip()), level=4)
            i += 1
            continue

        # ----- 구분선 -----
        if stripped == "---" or stripped == "***":
            doc.add_paragraph("")
            i += 1
            continue

        # ----- 글머리표 -----
        if stripped.startswith(("- ", "* ")):
            para = doc.add_paragraph(style="List Bullet")
            _add_inline_runs(para, stripped[2:].strip())
            i += 1
            continue

        # ----- 번호 매기기 (1. 2. 3.) -----
        import re
        m = re.match(r"^(\d+)\.\s+(.*)", stripped)
        if m:
            para = doc.add_paragraph(style="List Number")
            _add_inline_runs(para, m.group(2))
            i += 1
            continue

        # ----- 일반 단락 (인라인 마크업 처리) -----
        para = doc.add_paragraph()
        _add_inline_runs(para, line)
        i += 1


def _add_inline_runs(paragraph, text: str) -> None:
    """
    문자열 안의 **bold**, *italic*, [link](url) 등을
    여러 Run으로 분리하여 paragraph에 추가.
    """
    import re

    # 토큰 패턴: **bold** | *italic* | [text](url)
    # 순서 중요: **를 먼저 매칭 (그렇지 않으면 *italic*과 충돌)
    pattern = re.compile(
        r"(\*\*([^*]+)\*\*)"        # group 1,2: **bold**
        r"|(\*([^*]+)\*)"           # group 3,4: *italic*
        r"|(\[([^\]]+)\]\(([^)]+)\))"  # group 5,6,7: [text](url)
    )

    pos = 0
    for m in pattern.finditer(text):
        # 매칭 전 일반 텍스트
        if m.start() > pos:
            paragraph.add_run(text[pos:m.start()])

        if m.group(1):  # **bold**
            run = paragraph.add_run(m.group(2))
            run.bold = True
        elif m.group(3):  # *italic*
            run = paragraph.add_run(m.group(4))
            run.italic = True
        elif m.group(5):  # [text](url)
            link_text = m.group(6)
            url = m.group(7)
            # 단순 처리: 텍스트만 표시 + URL을 괄호로 (하이퍼링크 추가는 복잡)
            run = paragraph.add_run(f"{link_text}")
            run.underline = True
            run.font.color.rgb = None  # 기본 색상 유지
            # URL 별도 표시
            paragraph.add_run(f" ({url})")

        pos = m.end()

    # 남은 일반 텍스트
    if pos < len(text):
        paragraph.add_run(text[pos:])


def _strip_inline_marks(text: str) -> str:
    """헤딩 텍스트에서 인라인 마크업 제거 (간단히)."""
    import re
    # **bold** → bold
    text = re.sub(r"\*\*([^*]+)\*\*", r"\1", text)
    # *italic* → italic
    text = re.sub(r"\*([^*]+)\*", r"\1", text)
    # [text](url) → text
    text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)
    return text.strip()


def _add_code_block(doc, code: str) -> None:
    """코드 블록을 고정폭 폰트 단락으로 추가."""
    para = doc.add_paragraph()
    run = para.add_run(code)
    try:
        from docx.shared import Pt
        run.font.name = "Consolas"
        run.font.size = Pt(10)
    except Exception:
        pass


def _add_markdown_table(doc, table_lines: list) -> None:
    """
    마크다운 표 라인 목록을 docx Table로 변환.

    table_lines 예시:
      ['| Header1 | Header2 |',
       '|---------|---------|',
       '| Data1   | Data2   |',
       '| Data3   | Data4   |']
    """
    # 각 행을 셀로 분해
    rows = []
    for line in table_lines:
        line = line.strip()
        if not line.startswith("|"):
            continue
        # 양 끝 | 제거 후 분리
        cells = [c.strip() for c in line.strip("|").split("|")]
        rows.append(cells)

    if len(rows) < 2:
        return  # 헤더 + 구분선 최소 2줄 필요

    # 두 번째 줄(구분선) 제거
    header = rows[0]
    data_rows = rows[2:] if len(rows) > 2 else []

    # docx 표 생성
    n_cols = len(header)
    n_rows = 1 + len(data_rows)
    table = doc.add_table(rows=n_rows, cols=n_cols)
    table.style = "Light Grid Accent 1"  # 기본 격자 스타일

    # 헤더
    hdr_cells = table.rows[0].cells
    for j, cell_text in enumerate(header):
        if j < n_cols:
            hdr_cells[j].text = ""
            para = hdr_cells[j].paragraphs[0]
            run = para.add_run(_strip_inline_marks(cell_text))
            run.bold = True

    # 데이터
    for i, row in enumerate(data_rows):
        cells = table.rows[i + 1].cells
        for j, cell_text in enumerate(row):
            if j < n_cols:
                cells[j].text = ""
                para = cells[j].paragraphs[0]
                _add_inline_runs(para, _strip_inline_marks(cell_text))


# ---------------------------------------------------------------------------
# 헬스체크
# ---------------------------------------------------------------------------

@app.get("/health")
async def health():
    return {"status": "ok", "version": "4.0.0"}
